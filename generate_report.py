import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def add_toc(document):
    p = document.add_paragraph()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
    p._p.append(fld)


def add_field(paragraph, instruction):
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')
    paragraph._p.append(fld_char_begin)

    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = instruction
    paragraph._p.append(instr_text)

    fld_char_separate = OxmlElement('w:fldChar')
    fld_char_separate.set(qn('w:fldCharType'), 'separate')
    paragraph._p.append(fld_char_separate)

    run_result = OxmlElement('w:r')
    text = OxmlElement('w:t')
    text.text = ' '
    run_result.append(text)
    paragraph._p.append(run_result)

    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')
    paragraph._p.append(fld_char_end)


def set_header_footer(section):
    header = section.header
    header_p = header.paragraphs[0]
    header_p.text = 'Internship Report - Symfony Webapp (MJC)'
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.add_run('Page ')
    add_field(footer_p, 'PAGE')
    footer_p.add_run(' / ')
    add_field(footer_p, 'NUMPAGES')


def add_cover(document):
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Internship Report')
    run.bold = True
    run.font.size = Pt(28)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = subtitle.add_run('Symfony Web Application with EasyAdmin, Doctrine, Twig, VichUploader')
    r2.font.size = Pt(14)

    document.add_paragraph()

    info = document.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = info.add_run('Student: <Your Name>')
    r3.font.size = Pt(12)

    info2 = document.add_paragraph()
    info2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = info2.add_run('Company: <Company Name>  Mentor: <Mentor Name>')
    r4.font.size = Pt(12)

    info3 = document.add_paragraph()
    info3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = info3.add_run('Internship period: <Start Date> to <End Date>')
    r5.font.size = Pt(12)

    date_p = document.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r6 = date_p.add_run(datetime.date.today().strftime('%B %d, %Y'))
    r6.font.size = Pt(12)

    document.add_page_break()


def add_toc_page(document):
    h = document.add_paragraph('Table of Contents')
    h.style = document.styles['Heading 1']
    add_toc(document)
    document.add_page_break()


def add_sample_structure(document):
    for h1 in [
        'Introduction', 'Company Overview', 'Project Context', 'Technical Architecture',
        'Design and Technical Decisions', 'Implemented Features', 'Methodology and Workflow',
        'Testing and Performance', 'Challenges and Solutions', 'Security and Compliance',
        'Results and Impact', 'Limitations and Future Work', 'Conclusion', 'Appendices']:
        p = document.add_paragraph(h1)
        p.style = document.styles['Heading 1']
        document.add_paragraph('Placeholder content. Replace with your text.')
        for h2 in ['Overview', 'Details']:
            p2 = document.add_paragraph(h2)
            p2.style = document.styles['Heading 2']
            document.add_paragraph('Placeholder content. Replace with your text.')


def main():
    doc = Document()

    add_cover(doc)

    for section in doc.sections:
        set_header_footer(section)

    add_toc_page(doc)

    add_sample_structure(doc)

    doc.save('rapport_stage.docx')


if __name__ == '__main__':
    main()
